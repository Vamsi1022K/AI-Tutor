import shutil, os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Copy user's architecture image into project folder so Word can embed it
ARCH_SRC = r'C:\Users\vamsi\.gemini\antigravity\brain\1743dad4-7301-48a3-ad0b-dc032fa0b225\media__1776676875692.png'
ARCH_DEST = r'c:\Users\vamsi\OneDrive\Documents\Desktop\Sem 4\compiler design\cd ag\architecture_diagram.png'

# Only copy if the source file exists (user-provided architecture image)
if os.path.exists(ARCH_SRC):
    shutil.copy2(ARCH_SRC, ARCH_DEST)
    print(f"Architecture image copied to: {ARCH_DEST}")
else:
    print(f"WARNING: Architecture image not found at {ARCH_SRC}")
    ARCH_DEST = None

def set_cell_bold(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

def add_heading_with_style(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def main():
    doc = Document()

    # ----- Page margins -----
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ===========================
    # COVER PAGE
    # ===========================
    title = doc.add_heading('AI-Based Compiler Error Message Rewriting', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_title = doc.add_heading('AI Compiler Tutor', level=1)
    sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cover_info = doc.add_paragraph()
    cover_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_info.add_run(
        'Project-Based Learning (PBL) Report\n'
        'Compiler Design · Semester 4\n'
        'Department of Computer Science and Engineering\n'
        'National Institute of Technology, Warangal\n'
        'April 2026'
    ).font.size = Pt(13)

    doc.add_page_break()

    # ===========================
    # 1. ABSTRACT
    # ===========================
    add_heading_with_style(doc, '1. Abstract', level=1)
    doc.add_paragraph(
        "GCC error messages are designed for compiler engineers, not beginner students. "
        "This project—'AI Compiler Tutor'—intercepts GCC's cryptic output and transforms it into "
        "plain-English explanations with actionable code fixes and security-aware suggestions. "
        "It uses a hybrid architecture: deterministic rule-based matching handles 78% of errors, "
        "while a fine-tuned Google Flan-T5 neural model covers the remaining 22%. "
        "Evaluation shows a 222% improvement in clarity and 245% improvement in actionability "
        "over native GCC, while maintaining 100% diagnostic accuracy."
    )

    # ===========================
    # 2. PROBLEM STATEMENT
    # ===========================
    add_heading_with_style(doc, '2. Problem Statement', level=1)
    doc.add_paragraph(
        "When a beginner writes 'int x = 5' without a semicolon, GCC reports: "
        "'error: expected \';\' before \'printf\''. While technically accurate, "
        "this assumes knowledge of C's parser and token stream. The core problem spans "
        "four dimensions:"
    )
    problems = [
        "Lexical Opacity: Terms like 'lvalue', 'rvalue', and 'implicit declaration' are meaningless to novices.",
        "Lack of Actionability: Errors describe what is wrong, not how to fix it.",
        "Security Blindness: GCC does not warn about dangerous but syntactically valid patterns (e.g., gets()).",
        "Cascade Effects: A single mistake triggers multiple errors, overwhelming beginners.",
    ]
    for p in problems:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_paragraph(
        "The research gap: no system transparently bridges standard GCC workflows with beginner-friendly, "
        "security-integrated, offline-capable explanations. AI Compiler Tutor fills that gap."
    )

    # ===========================
    # 3. OBJECTIVES
    # ===========================
    add_heading_with_style(doc, '3. Objectives', level=1)
    objectives = [
        "Intercept & Parse: Capture raw GCC stderr and parse into structured error data.",
        "Error Rewriting: Replace opaque jargon with plain English explanations.",
        "Actionable Recommendations: Provide concrete, copy-paste code fixes.",
        "Security-Aware Suggestions: Block unsafe patterns (gets, strcpy) and suggest secure alternatives.",
        "AI Fallback: Use a fine-tuned Flan-T5 model for novel or complex errors.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # ===========================
    # 4. SYSTEM ARCHITECTURE
    # ===========================
    doc.add_page_break()
    add_heading_with_style(doc, '4. System Architecture', level=1)
    doc.add_paragraph(
        "The AI Compiler Tutor operates as a transparent intermediary between the programmer and GCC. "
        "Figure 1 below shows the complete processing pipeline."
    )

    # Embed user's architecture diagram
    if ARCH_DEST and os.path.exists(ARCH_DEST):
        try:
            doc.add_picture(ARCH_DEST, width=Inches(5.8))
            cap = doc.add_paragraph('Figure 1: AI Compiler Tutor — System Architecture')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(10)
        except Exception as e:
            doc.add_paragraph(f'[Could not embed architecture image: {e}]')
    else:
        doc.add_paragraph('[Architecture diagram not available]')

    doc.add_paragraph('')

    doc.add_paragraph(
        "The six core processing stages (as depicted in Figure 1):"
    )
    stages = [
        ("Input Handling Module", "Accepts the .c source file from the user and validates it before passing to the compiler wrapper."),
        ("Compiler Wrapper → System GCC", "Invokes gcc -Wall -fsyntax-only as a sub-process, capturing stderr/exit code without generating a binary."),
        ("Error Parser Engine", "Converts raw stderr text into structured error objects: file, line, column, error type, and message."),
        ("Rewriting & Logic Module", "The hybrid engine — queries the Explanation Database (rule-based, 115 patterns) and Source Code Reader (AST context) before invoking the AI model."),
        ("Security Filter", "Intercepts all fix suggestions and blocks unsafe patterns (gets, strcpy, sprintf, unbounded scanf). Produces verified, safe output."),
        ("Output Formatter → User", "Presents the beginner-friendly explanation with corrected code back to the user via CLI, Web, or Desktop GUI."),
    ]
    for title_s, desc in stages:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title_s + ": ").bold = True
        p.add_run(desc)

    # ===========================
    # 5. IMPLEMENTATION
    # ===========================
    doc.add_page_break()
    add_heading_with_style(doc, '5. Implementation', level=1)

    add_heading_with_style(doc, '5.1 Technology Stack', level=2)
    doc.add_paragraph(
        "Python 3.9+, PyTorch 2.0, Hugging Face Transformers 4.30 (Flan-T5-base, 250M params), "
        "pycparser 2.21 for AST construction, Flask 2.3 for the web interface, and tkinter for the desktop GUI."
    )

    add_heading_with_style(doc, '5.2 Dataset & Model Training', level=2)
    doc.add_paragraph(
        "Errors are classified into 5 GCC pipeline stages: Syntax, Declaration, Type, Scope, and Linker. "
        "Starting from 10 base examples per category, systematic augmentation (identifier substitution, "
        "literal variation, comment injection) generated 11,550 training examples. "
        "Flan-T5-base was fine-tuned for 3 epochs (lr=5e-5, batch=16) achieving BLEU-4 = 0.92 on held-out test data. "
        "Rule-based matching (115 patterns) handles 78% of errors in under 5ms; the AI model covers the remaining 22%."
    )

    add_heading_with_style(doc, '5.3 Security Integration', level=2)
    doc.add_paragraph(
        "secure_checker.py maintains a blocklist of dangerous functions: gets(), strcpy(), strcat(), "
        "sprintf() (without bounds), and unbounded scanf(). It runs at two points: "
        "(1) proactively scanning source before compilation, and (2) validating all AI-suggested fixes. "
        "This embeds secure coding habits from students' first exposure to C."
    )

    # ===========================
    # 6. EVALUATION
    # ===========================
    add_heading_with_style(doc, '6. Evaluation & Results', level=1)
    doc.add_paragraph(
        "Evaluated on 20 representative C programs using a 1–5 scale rubric scored by two independent reviewers "
        "(Cohen's kappa = 0.89):"
    )

    # Results table
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    headers = ['Metric', 'GCC Default', 'AI Tutor', 'Improvement']
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bold(hdr[i])

    results = [
        ('Clarity',       '1.55 / 5.0', '5.0 / 5.0', '+222%'),
        ('Actionability', '1.45 / 5.0', '5.0 / 5.0', '+245%'),
        ('Accuracy',      '5.0 / 5.0',  '5.0 / 5.0', 'Maintained'),
        ('Security',      '3.0 / 5.0',  '5.0 / 5.0', '+67%'),
    ]
    for row_data in results:
        row = tbl.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph('')
    doc.add_paragraph(
        "Error processing averages 0.23 seconds per error on a modern CPU. Rule-based matching executes "
        "in under 5ms; neural inference (100–200ms) occurs for only 22% of cases. Combined error coverage "
        "effectively reaches 100%."
    )

    # ===========================
    # 7. CONCLUSION
    # ===========================
    add_heading_with_style(doc, '7. Conclusion & Future Work', level=1)
    doc.add_paragraph(
        "The AI Compiler Tutor demonstrates that compiler tooling can bridge technical precision and "
        "pedagogical accessibility. The hybrid architecture achieves 100% error coverage while maintaining "
        "GCC's diagnostic accuracy. Security education is embedded directly in the feedback loop, instilling "
        "safe habits from the start."
    )
    doc.add_paragraph("Key future directions include:")
    future = [
        "C++ Template Error Handling: Summarizing complex template instantiation failures.",
        "Interactive Dialogue: Conversational follow-up questions for deeper pedagogical value.",
        "Personalization: Adapting explanations to individual student knowledge levels.",
        "IDE Plugins: VS Code / CLion integration for seamless inline assistance.",
        "Runtime Errors: Expanding beyond compilation to segfaults and undefined behavior.",
    ]
    for f in future:
        doc.add_paragraph(f, style='List Bullet')

    out_path = r'c:\Users\vamsi\OneDrive\Documents\Desktop\Sem 4\compiler design\cd ag\AI_Compiler_Tutor_PBL_Report_Reduced.docx'
    doc.save(out_path)
    print(f"Report saved: {out_path}")
    print("Estimated reading time: ~5 minutes")

if __name__ == "__main__":
    main()
