from docx import Document
from docx.shared import Pt, Inches

def main():
    doc_path = r'c:\Users\vamsi\OneDrive\Documents\Desktop\Sem 4\compiler design\cd ag\AI_Compiler_Tutor_PBL_Report.docx'
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error opening doc: {e}")
        return

    doc.add_page_break()
    p_main = doc.add_paragraph()
    r_main = p_main.add_run('System Architecture')
    r_main.bold = True
    r_main.font.size = Pt(20)
    
    doc.add_paragraph(
        "The AI Compiler Tutor intercepts the standard C compilation process to dynamically "
        "analyze and rewrite cryptic GCC errors into beginner-friendly explanations with verified source code fixes."
    )

    p_wf = doc.add_paragraph()
    r_wf = p_wf.add_run('Pipeline Workflow')
    r_wf.bold = True
    r_wf.font.size = Pt(16)
    
    workflow = [
        ("Step 1", "GCC Interception", "The system programmatically invokes gcc -Wall -fsyntax-only, capturing the raw stderr output without generating binary artifacts."),
        ("Step 2", "Error Parsing", "The raw stderr is digested via Regular Expressions into structured segments: file path, line number, column, and explicit error message."),
        ("Step 3", "AST Context Retrieval", "ast_parser.py constructs an Abstract Syntax Tree to pinpoint the exact semantic construct (e.g., VarDecl, CallExpr) causing the fault."),
        ("Step 4", "Hybrid Matcher Engine", "Before invoking neural models, the system queries annotated_errors.json. If a static structural pattern matches, it retrieves highly-curated plain text explanations."),
        ("Step 5", "Flan-T5 Fallback Inference", "For unrecognized or novel compilation errors, a local PyTorch instance of Google Flan-T5 infers a custom fix using contextual prompts."),
        ("Step 6", "Security Guardrails", "All suggested code-fixes pass through secure_checker.py to block unsafe paradigms like gets() or non-bound string copies."),
    ]
    
    for step, title, desc in workflow:
        p = doc.add_paragraph()
        run1 = p.add_run(f"[{step}] {title}: ")
        run1.bold = True
        p.add_run(desc)
        
    p_comp = doc.add_paragraph()
    r_comp = p_comp.add_run('\nComponent Modules\n')
    r_comp.bold = True
    r_comp.font.size = Pt(16)

    modules = [
        ("server.py", "Flask REST API bridge & GCC sub-process execution."),
        ("main.py", "Standalone CLI interface for the entire AI pipeline."),
        ("secure_checker.py", "Static AST logic restricting unsecure memory calls."),
        ("evaluate.py", "Generates benchmark scores on Clarity, Security, and Accuracy."),
    ]
    for comp, resp in modules:
        p = doc.add_paragraph()
        run = p.add_run(comp + ": ")
        run.bold = True
        p.add_run(resp)

    doc.save(doc_path)
    print("Successfully appended Architecture to the report!")

if __name__ == '__main__':
    main()
